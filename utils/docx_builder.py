"""
utils/docx_builder.py — Convert resume / cover-letter text to a formatted DOCX file.
"""

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _add_runs_with_bold(paragraph, text: str):
    """Add text runs to a paragraph, honouring **bold** markers."""
    if not text:
        return
    last = 0
    found = False
    for m in BOLD_PATTERN.finditer(text):
        found = True
        pre = text[last : m.start()]
        if pre:
            paragraph.add_run(pre)
        paragraph.add_run(m.group(1)).bold = True
        last = m.end()
    remaining = text[last:]
    if not found:
        remaining = re.sub(r"\*\*", "", remaining)
    if remaining:
        paragraph.add_run(remaining)


def _add_horizontal_line(doc):
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_docx_from_text(text: str, name: str = "") -> bytes:
    """
    Convert plain / lightly-marked-up resume text into a polished DOCX.

    Returns raw bytes suitable for st.download_button.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.15
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    if name and not text.strip().startswith(name):
        h = doc.add_paragraph()
        r = h.add_run(name.strip())
        r.bold = True
        r.font.size = Pt(18)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.space_after = Pt(6)
        doc.add_paragraph()

    lines = text.splitlines()
    i = 0
    in_core_skills = False

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Section header
        if (line.isupper() and len(line.split()) <= 4) or re.match(
            r"^(PROFESSIONAL SUMMARY|CORE SKILLS|PROFESSIONAL EXPERIENCE|"
            r"WORK EXPERIENCE|EDUCATION|CERTIFICATIONS?|TECHNICAL SKILLS|PROJECTS)[\s:]*$",
            line,
            re.I,
        ):
            in_core_skills = bool(
                re.match(r"^(CORE SKILLS|TECHNICAL SKILLS)[\s:]*$", line, re.I)
            )
            if i > 0:
                _add_horizontal_line(doc)

            heading_text = re.sub(r"[#\*_\-]{2,}", "", line.strip().rstrip(":")).strip()
            h = doc.add_heading(heading_text, level=1)
            h.runs[0].font.size = Pt(14)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # Skill category line (no bullets)
        if (
            ("–" in line or "—" in line or (":" in line and len(line.split(":")[0].split()) <= 6))
            and not line.startswith("[")
            and "|" not in line
        ):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)

            if "–" in line:
                cat, *rest = line.split("–", 1)
                sep = " – "
            elif "—" in line:
                cat, *rest = line.split("—", 1)
                sep = " — "
            else:
                cat, *rest = line.split(":", 1)
                sep = ": "

            p.add_run(re.sub(r"\*\*", "", cat.strip())).bold = True
            if rest:
                p.add_run(sep + re.sub(r"\*\*", "", rest[0].strip()))
            i += 1
            continue

        # Bullet point
        if not in_core_skills and re.match(r"^[\•\-\*]\s+", line):
            bullet_text = re.sub(r"^[\•\-\*]\s+", "", line).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            _add_runs_with_bold(p, bullet_text)
            i += 1
            continue

        # Job title / pipe line — entire line bold
        if "|" in line and not line.startswith("["):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(re.sub(r"\*\*", "", line.strip())).bold = True
            i += 1
            continue

        # Contact info — centred
        if any(kw in line.lower() for kw in ["@", "phone", "email", "linkedin", "github", "number"]):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)
            _add_runs_with_bold(p, line)
            i += 1
            continue

        # Generic paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        _add_runs_with_bold(p, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
